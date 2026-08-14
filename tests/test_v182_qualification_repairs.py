import json
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "scripts"))

import router
from common_utils import markdown_link_label


def test_irregular_csv_preserves_extra_field_and_warns(tmp_path):
    source = tmp_path / "irregular.csv"
    source.write_text("a,b,c\ntwo,columns,but_file_has_more,extra\nshort,row\n", encoding="utf-8")
    bundle = tmp_path / "bundle"
    report = router.convert(str(source), str(bundle))
    markdown = (bundle / "document.md").read_text(encoding="utf-8")
    assert report["status"] == "passed_with_warnings"
    assert "CSV_INCONSISTENT_COLUMN_COUNT" in [warning["code"] for warning in report["warnings"]]
    assert "__extra_1" in markdown and "extra" in markdown


def test_crlf_csv_uses_comma_delimiter(tmp_path):
    source = tmp_path / "crlf.csv"
    source.write_bytes(b"name,note\r\nalpha,ordinary\r\nbeta,final\r\n")
    bundle = tmp_path / "bundle"

    report = router.convert(str(source), str(bundle))
    persisted = json.loads((bundle / "conversion-report.json").read_text(encoding="utf-8"))

    assert report["status"] == "passed"
    assert persisted["details"]["max_column_count"] == 2
    assert "| alpha | ordinary |" in (bundle / "document.md").read_text(encoding="utf-8")


def test_crlf_quoted_multiline_csv_preserves_extra_columns_and_warning(tmp_path):
    source = tmp_path / "crlf-multiline.csv"
    source.write_bytes(
        b'name,note\r\nalpha,"line 1\r\nline 2"\r\nbeta,=SUM(1,1),EXTRA-A,EXTRA-B\r\n'
    )
    bundle = tmp_path / "bundle"

    report = router.convert(str(source), str(bundle))
    persisted = json.loads((bundle / "conversion-report.json").read_text(encoding="utf-8"))
    markdown = (bundle / "document.md").read_text(encoding="utf-8")

    assert report["status"] == "passed_with_warnings"
    assert persisted["details"]["max_column_count"] == 5
    assert "CSV_INCONSISTENT_COLUMN_COUNT" in [warning["code"] for warning in persisted["warnings"]]
    assert all(label in markdown for label in ("__extra_1", "__extra_2", "__extra_3", "EXTRA-A", "EXTRA-B"))
    assert "line 1\n<br>line 2" in markdown


def test_lf_quoted_multiline_csv_matches_crlf_canonical_grid(tmp_path):
    crlf_source = tmp_path / "crlf.csv"
    lf_source = tmp_path / "lf.csv"
    crlf_source.write_bytes(
        b'name,note\r\nalpha,"line 1\r\nline 2"\r\nbeta,=SUM(1,1),EXTRA-A,EXTRA-B\r\n'
    )
    lf_source.write_text(
        'name,note\nalpha,"line 1\nline 2"\nbeta,=SUM(1,1),EXTRA-A,EXTRA-B\n', encoding="utf-8"
    )
    crlf_bundle, lf_bundle = tmp_path / "crlf-bundle", tmp_path / "lf-bundle"

    router.convert(str(crlf_source), str(crlf_bundle))
    router.convert(str(lf_source), str(lf_bundle))
    crlf_report = json.loads((crlf_bundle / "conversion-report.json").read_text(encoding="utf-8"))
    lf_report = json.loads((lf_bundle / "conversion-report.json").read_text(encoding="utf-8"))

    assert crlf_report["details"]["max_column_count"] == lf_report["details"]["max_column_count"] == 5
    assert crlf_report["warnings"] == lf_report["warnings"]
    assert crlf_report["status"] == lf_report["status"] == "passed_with_warnings"


def test_json_nesting_limit_has_stable_failure_reason(tmp_path):
    value = "leaf"
    for _ in range(1100):
        value = {"a": value}
    source = tmp_path / "deep.json"
    source.write_text(json.dumps(value), encoding="utf-8")
    bundle = tmp_path / "bundle"
    report = router.convert(str(source), str(bundle))
    assert report["status"] == "failed"
    assert report["reason"] == "JSON_NESTING_LIMIT_EXCEEDED"
    assert report["maximum_depth"] > report["configured_limit"]
    persisted = json.loads((bundle / "conversion-report.json").read_text(encoding="utf-8"))
    assert persisted["reason"] == "JSON_NESTING_LIMIT_EXCEEDED"


def test_json_extreme_nesting_is_rejected_before_json_loads_recurses(tmp_path):
    source = tmp_path / "very-deep.json"
    source.write_text('{"a":' * 10_000 + '"leaf"' + '}' * 10_000, encoding="utf-8")
    bundle = tmp_path / "bundle"
    router.convert(str(source), str(bundle))
    persisted = json.loads((bundle / "conversion-report.json").read_text(encoding="utf-8"))
    assert persisted["status"] == "failed"
    assert persisted["reason"] == "JSON_NESTING_LIMIT_EXCEEDED"
    assert persisted["maximum_depth"] == 10_000


def test_email_attachment_is_visible_in_readable_projection_and_chunk_has_source_file(tmp_path):
    source = tmp_path / "message.eml"
    source.write_text(
        "From: sender@example.com\nTo: recipient@example.com\nSubject: test\n"
        "MIME-Version: 1.0\nContent-Type: multipart/mixed; boundary=boundary\n\n"
        "--boundary\nContent-Type: text/plain\n\nBody\n"
        "--boundary\nContent-Type: text/plain; name=report.txt\n"
        "Content-Disposition: attachment; filename=report.txt\n"
        "Content-Transfer-Encoding: base64\n\nY29udGVudAo=\n--boundary--\n",
        encoding="utf-8",
    )
    bundle = tmp_path / "bundle"
    report = router.convert(str(source), str(bundle))
    assert report["status"] == "passed"
    assert "[report.txt](assets/report.txt)" in (bundle / "document.md").read_text(encoding="utf-8")
    chunks = [json.loads(line) for line in (bundle / "chunks.jsonl").read_text(encoding="utf-8").splitlines()]
    assert chunks and all(chunk["source_file"] == "message.eml" for chunk in chunks)


def test_email_attachment_name_cannot_inject_markdown_structure():
    label = markdown_link_label("bad](https://attacker.invalid)\n# heading.txt")
    assert label == "bad\\](https://attacker.invalid) # heading.txt"


def test_digital_pdf_with_material_raster_is_not_silent_success(tmp_path):
    import fitz
    from PIL import Image, ImageDraw, ImageFont

    image = tmp_path / "scan.png"
    raster = Image.new("RGB", (600, 800), "white")
    ImageDraw.Draw(raster).text((40, 80), "PUMP TAG: P-204", fill="black", font=ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 36))
    raster.save(image)
    source = tmp_path / "hybrid.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((72, 50), "Scanned page below")
    page.insert_image(fitz.Rect(40, 80, 555, 760), filename=str(image))
    pdf.save(source)
    bundle = tmp_path / "bundle"
    report = router.convert(str(source), str(bundle))
    assert report["status"] == "passed_with_warnings"
    assert report["details"]["ocr_used"] is True
    assert "MIXED_PDF_MATERIAL_RASTER_OCR" in [warning["code"] for warning in report["warnings"]]
    assert "P-204" in (bundle / "document.md").read_text(encoding="utf-8")


def test_true_mixed_pdf_ocrs_only_material_raster_region(tmp_path):
    import fitz
    from PIL import Image, ImageDraw, ImageFont

    image = tmp_path / "scan.png"
    raster = Image.new("RGB", (600, 500), "white")
    ImageDraw.Draw(raster).text((40, 80), "PUMP TAG: P-204", fill="black", font=ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", 36))
    raster.save(image)
    source = tmp_path / "true-mixed.pdf"; pdf = fitz.open()
    scanned = pdf.new_page(); scanned.insert_image(fitz.Rect(30, 30, 565, 800), filename=str(image))
    hybrid = pdf.new_page(); hybrid.insert_text((72, 50), "Native heading must appear once")
    hybrid.insert_image(fitz.Rect(40, 100, 555, 650), filename=str(image))
    pdf.save(source)
    bundle = tmp_path / "bundle"
    router.convert(str(source), str(bundle))
    report = json.loads((bundle / "conversion-report.json").read_text(encoding="utf-8"))
    markdown = (bundle / "document.md").read_text(encoding="utf-8")
    assert report["details"]["ocr_used"] is True
    assert report["details"]["hybrid_raster_pages"] == [2]
    assert markdown.count("Native heading must appear once") == 1
    assert markdown.count("P-204") >= 2


def test_nested_docx_table_is_explicitly_delimited_and_warned(tmp_path):
    from docx import Document
    source = tmp_path / "nested.docx"
    document = Document(); cell = document.add_table(rows=1, cols=1).cell(0, 0)
    nested = cell.add_table(rows=2, cols=2)
    for row in range(2):
        for column in range(2): nested.cell(row, column).text = f"N{row}{column}"
    document.save(source)
    bundle = tmp_path / "bundle"
    report = router.convert(str(source), str(bundle))
    assert report["status"] == "passed_with_warnings"
    assert "DOCX_NESTED_TABLE_FLATTENED" in [warning["code"] for warning in report["warnings"]]
    assert "Nested table: N00 \\| N01 \\| N10 \\| N11" in (bundle / "document.md").read_text(encoding="utf-8")


def test_nested_docx_table_preserves_outer_cell_text_in_document_order(tmp_path):
    from docx import Document

    source = tmp_path / "nested-with-outer-text.docx"
    document = Document(); cell = document.add_table(rows=1, cols=1).cell(0, 0)
    cell.text = "OUTER-INTRO"
    nested = cell.add_table(rows=1, cols=2)
    nested.cell(0, 0).text = "INNER-A"; nested.cell(0, 1).text = "INNER-B"
    cell.add_paragraph("OUTER-TAIL")
    document.save(source)
    bundle = tmp_path / "bundle"
    router.convert(str(source), str(bundle))
    markdown = (bundle / "document.md").read_text(encoding="utf-8")
    assert "OUTER-INTRO" in markdown and "OUTER-TAIL" in markdown
    assert "Nested table: INNER-A \\| INNER-B" in markdown
    assert markdown.index("OUTER-INTRO") < markdown.index("Nested table") < markdown.index("OUTER-TAIL")
