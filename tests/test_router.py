"""
test_router.py
Regression tests for the bugs found and fixed during this skill's
development. Each test is named after the behavior it locks in, not just
"it doesn't crash" - a passing test should mean the specific fix still
works, so a future refactor of _render_table() or the OCR heuristics
can't silently regress it.
"""

import json
import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "scripts"))

import router


def _run(input_path, tmp_out):
    return router.convert(input_path, tmp_out)


class TestXlsxMergedCells:
    def test_merged_range_becomes_html_colspan(self, xlsx_merged, tmp_out):
        report = _run(xlsx_merged, tmp_out)
        assert report["status"] == "passed"
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert 'colspan="2"' in md
        assert "Header" in md

    def test_plain_sheet_stays_pipe_table_not_html(self, tmp_path, tmp_out):
        import openpyxl
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["A", "B"])
        ws.append([1, 2])
        path = str(tmp_path / "plain.xlsx")
        wb.save(path)
        _run(path, tmp_out)
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert "<table>" not in md
        assert "| A | B |" in md


class TestXlsxFormulaPreservation:
    def test_missing_cached_value_is_not_silent_blank(self, xlsx_formula_no_cache, tmp_out):
        report = _run(xlsx_formula_no_cache, tmp_out)
        assert report["status"] == "passed_with_warnings"
        codes = [w["code"] for w in report["warnings"]]
        assert "FORMULA_RESULT_UNAVAILABLE" in codes
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert "=A3+B3" in md  # formula preserved, not an empty cell


class TestDocxMergedCells:
    def test_horizontal_merge_colspan(self, docx_horizontal_merge, tmp_out):
        report = _run(docx_horizontal_merge, tmp_out)
        assert report["status"] == "passed"
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert 'colspan="2"' in md
        assert "Merged Header" in md

    def test_vertical_merge_rowspan(self, docx_vertical_merge, tmp_out):
        report = _run(docx_vertical_merge, tmp_out)
        assert report["status"] == "passed"
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert 'rowspan="2"' in md
        assert "VMerge" in md


class TestDocxFidelity:
    def test_bold_and_italic_runs_preserved(self, docx_bold_italic, tmp_out):
        report = _run(docx_bold_italic, tmp_out)
        assert report["status"] == "passed"
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert "**bold text**" in md
        assert "*italic*" in md


class TestPdfClassification:
    def test_short_digital_pdf_not_misrouted_to_ocr(self, pdf_short_digital, tmp_out):
        """Regression for the exact P0 bug: a short digital PDF (well under
        the old 50-char whole-document threshold) must be classified
        digital on a per-page basis, not sent through OCR."""
        report = _run(pdf_short_digital, tmp_out)
        assert report["status"] == "passed"
        assert report["details"]["engine"] == "pdfplumber"
        assert report["details"]["ocr_used"] is False


class TestEncodingDetection:
    def test_big5_csv_decodes_correctly_not_utf16(self, csv_big5, tmp_out):
        """Regression for the exact P0 bug: a Big5-encoded CSV was
        misdetected as UTF-16BE, producing legal-but-wrong Unicode with
        status: passed and no warning."""
        report = _run(csv_big5, tmp_out)
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert "鋼管" in md
        assert report["details"]["encoding_used"] in ("big5", "cp950")

    def test_plain_utf8_csv_has_no_ambiguity_warning(self, tmp_path, tmp_out):
        path = str(tmp_path / "utf8.csv")
        with open(path, "w", encoding="utf-8") as f:
            f.write("name,score\nAlice,90\n")
        report = _run(path, tmp_out)
        assert report["status"] == "passed"
        assert report["details"]["encoding_ambiguous"] is False

    def test_explicit_encoding_resolves_short_cjk_ambiguity(self, tmp_path, tmp_out):
        path = str(tmp_path / "simplified.csv")
        with open(path, "wb") as f:
            f.write("名称,数量\n钢管,10\n".encode("gb18030"))
        report = router.convert(path, tmp_out, encoding_hint="gb18030")
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert report["status"] == "passed"
        assert "名称" in md and "钢管" in md
        assert report["details"]["encoding_user_selected"] is True

    def test_invalid_json_is_not_silent_success(self, tmp_path, tmp_out):
        path = str(tmp_path / "broken.json")
        with open(path, "w", encoding="utf-8") as f:
            f.write('{"missing": }')
        report = _run(path, tmp_out)
        assert report["status"] == "passed_with_warnings"
        assert "INVALID_JSON_PRESERVED" in [warning["code"] for warning in report["warnings"]]


class TestFormatDetection:
    def test_extension_mismatch_uses_real_content_type(self, xlsx_mislabeled_as_pdf, tmp_out):
        """An .xlsx file renamed to .pdf must still be converted as xlsx
        (via magic-byte/container detection), with the mismatch reported,
        not silently routed to the PDF converter or failing outright."""
        report = _run(xlsx_mislabeled_as_pdf, tmp_out)
        assert report["file_type"] == "xlsx"
        codes = [w["code"] for w in report["warnings"]]
        assert "FORMAT_EXTENSION_MISMATCH" in codes


class TestPptx:
    def test_merged_table_cell_becomes_colspan(self, pptx_merge_table, tmp_out):
        report = _run(pptx_merge_table, tmp_out)
        assert report["status"] == "passed"
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert 'colspan="2"' in md


class TestPptxBulletsV151:
    """Regression for the v1.5.1 P1 fix: level 0 is PowerPoint's ordinary
    top-level bullet indentation, not "not a bullet". The old code only
    rendered `- ` for level > 0, so a flat bulleted list (the most common
    case in real slides) lost its list semantics entirely."""

    def test_level_zero_paragraphs_render_as_bullets(self, pptx_level0_bullets, tmp_out):
        report = _run(pptx_level0_bullets, tmp_out)
        assert report["status"] == "passed"
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert "- First bullet point" in md
        assert "- Second bullet point" in md

    def test_explicit_buNone_paragraph_is_not_a_bullet(self, pptx_level0_bullets, tmp_out):
        _run(pptx_level0_bullets, tmp_out)
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert "- Not a bullet line" not in md
        assert "Not a bullet line" in md

    def test_plain_textbox_is_not_bullet(self, pptx_plain_textbox, tmp_out):
        _run(pptx_plain_textbox, tmp_out)
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert "Plain paragraph one" in md
        assert "- Plain paragraph one" not in md
        assert "- Plain paragraph two" not in md

    def test_explicit_buchar_and_autonum(self, pptx_explicit_bullets, tmp_out):
        _run(pptx_explicit_bullets, tmp_out)
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert "- Explicit char" in md
        assert "1. Explicit number" in md


class TestPptxGroupNestedTableV151:
    """Regression for the v1.5.1 P1 fix: a table nested inside a group
    shape used to render into document.md's Markdown but was silently
    discarded before reaching tables_out, so its standalone
    tables/*.csv+*.html asset never existed."""

    def test_table_inside_group_is_exported_as_standalone_asset(
            self, pptx_group_with_table, tmp_out):
        report = _run(pptx_group_with_table, tmp_out)
        assert report["status"] == "passed"
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert "GroupCell" in md
        tables_dir = os.path.join(tmp_out, "tables")
        assert os.path.isdir(tables_dir), "group-nested table produced no standalone asset"
        csvs = [f for f in os.listdir(tables_dir) if f.endswith(".csv")]
        assert csvs, "expected at least one exported table CSV for the group-nested table"
        content = open(os.path.join(tables_dir, csvs[0]), encoding="utf-8").read()
        assert "GroupCell" in content


class TestPptxSmartArtOleDisclosureV151:
    """Regression for the v1.5.1 P1 fix: SmartArt/OLE presence must be
    detected and surfaced in conversion-report.json, not just documented
    as an abstract known-limitation that no real file's report reflects."""

    def test_smartart_and_ole_relationships_produce_warnings(
            self, pptx_with_smartart_and_ole, tmp_out):
        report = _run(pptx_with_smartart_and_ole, tmp_out)
        assert report["status"] == "passed_with_warnings"
        codes = [w["code"] for w in report["warnings"]]
        assert "SMARTART_NOT_EXTRACTED" in codes
        assert "EMBEDDED_OLE_NOT_EXTRACTED" in codes
        details = report["details"]
        assert details["smartart_occurrences"][0]["slide"] == 1
        assert details["smartart_occurrences"][0]["relationship_id"] == "rId100"
        assert details["ole_occurrences"][0]["slide"] == 1

    def test_plain_pptx_has_no_smartart_ole_warnings(self, pptx_merge_table, tmp_out):
        report = _run(pptx_merge_table, tmp_out)
        codes = [w["code"] for w in report["warnings"]]
        assert "SMARTART_NOT_EXTRACTED" not in codes
        assert "EMBEDDED_OLE_NOT_EXTRACTED" not in codes


class TestPptxMergedCellHtmlSpanV151:
    """Regression for the v1.5.1 P1 fix: document.md preserved a merged
    cell's rowspan/colspan, but the exact same table's standalone
    tables/<id>.html silently flattened it back into a plain grid because
    table_export.py only ever received a flat `rows` grid."""

    def test_standalone_html_preserves_colspan(self, pptx_merge_table, tmp_out):
        _run(pptx_merge_table, tmp_out)
        tables_dir = os.path.join(tmp_out, "tables")
        html_files = [f for f in os.listdir(tables_dir) if f.endswith(".html")]
        assert html_files
        content = open(os.path.join(tables_dir, html_files[0]), encoding="utf-8").read()
        assert 'colspan="2"' in content, "standalone HTML lost merge geometry"


class TestDocxMergedCellHtmlSpanV151:
    def test_standalone_html_preserves_colspan(self, docx_horizontal_merge, tmp_out):
        _run(docx_horizontal_merge, tmp_out)
        tables_dir = os.path.join(tmp_out, "tables")
        html_files = [f for f in os.listdir(tables_dir) if f.endswith(".html")]
        assert html_files
        content = open(os.path.join(tables_dir, html_files[0]), encoding="utf-8").read()
        assert 'colspan="2"' in content


class TestXlsxMergedCellHtmlSpanV151:
    def test_standalone_html_preserves_colspan(self, xlsx_merged, tmp_out):
        _run(xlsx_merged, tmp_out)
        tables_dir = os.path.join(tmp_out, "tables")
        html_files = [f for f in os.listdir(tables_dir) if f.endswith(".html")]
        assert html_files
        content = open(os.path.join(tables_dir, html_files[0]), encoding="utf-8").read()
        assert 'colspan="2"' in content

    def test_csv_grid_is_rectangular_not_ragged(self, xlsx_merged, tmp_out):
        """The v1.5 grid passed to table_export.py silently dropped
        spanned-over cells instead of leaving a blank placeholder,
        producing a ragged (non-rectangular) CSV. Every row must have
        the same column count."""
        _run(xlsx_merged, tmp_out)
        tables_dir = os.path.join(tmp_out, "tables")
        csv_files = [f for f in os.listdir(tables_dir) if f.endswith(".csv")]
        assert csv_files
        import csv as csv_module
        with open(os.path.join(tables_dir, csv_files[0]), encoding="utf-8") as f:
            rows = list(csv_module.reader(f))
        widths = {len(r) for r in rows}
        assert len(widths) == 1, f"ragged CSV rows: widths found = {widths}"


class TestDocumentJsonElementSchemaV151:
    """Regression for the v1.5.1 P1 fix: document.json elements must all
    expose the same top-level keys (engine/confidence/source_locator),
    with a converter-level default filled in when a specific element
    didn't set its own engine."""

    def test_docx_elements_have_normalized_schema(self, docx_horizontal_merge, tmp_out):
        _run(docx_horizontal_merge, tmp_out)
        doc_json = json.load(open(os.path.join(tmp_out, "document.json"), encoding="utf-8"))
        for el in doc_json["elements"]:
            assert "engine" in el
            assert "confidence" in el
            assert "source_locator" in el
            assert el["engine"] is not None

    def test_pptx_elements_have_normalized_schema(self, pptx_merge_table, tmp_out):
        _run(pptx_merge_table, tmp_out)
        doc_json = json.load(open(os.path.join(tmp_out, "document.json"), encoding="utf-8"))
        for el in doc_json["elements"]:
            assert "engine" in el
            assert "confidence" in el
            assert "source_locator" in el


class TestPdfTableLikelihoodV151:
    """Regression for the v1.5.1 P1 fix: TABLE_STRUCTURE_UNVERIFIED must
    only fire when the page plausibly looked tabular (table_likelihood),
    not as a blanket disclaimer on every scanned page with zero detected
    tables regardless of actual content."""

    def test_column_aligned_page_triggers_warning(self, pdf_column_aligned_scanned, tmp_out):
        report = _run(pdf_column_aligned_scanned, tmp_out)
        codes = [w["code"] for w in report["warnings"]]
        assert "TABLE_STRUCTURE_UNVERIFIED" in codes

    def test_plain_prose_page_does_not_trigger_warning(self, pdf_plain_scanned_no_table, tmp_out):
        report = _run(pdf_plain_scanned_no_table, tmp_out)
        codes = [w["code"] for w in report["warnings"]]
        assert "TABLE_STRUCTURE_UNVERIFIED" not in codes


class TestOutputBundle:
    def test_bundle_has_all_v13_files(self, xlsx_merged, tmp_out):
        _run(xlsx_merged, tmp_out)
        for fname in ("document.md", "document.json", "chunks.jsonl",
                      "manifest.json", "conversion-report.json"):
            assert os.path.exists(os.path.join(tmp_out, fname)), f"missing {fname}"

    def test_document_json_has_elements(self, xlsx_merged, tmp_out):
        _run(xlsx_merged, tmp_out)
        doc_json = json.load(open(os.path.join(tmp_out, "document.json"), encoding="utf-8"))
        assert doc_json["element_count"] >= 1

    def test_tables_exported_as_csv_and_html(self, xlsx_merged, tmp_out):
        _run(xlsx_merged, tmp_out)
        tables_dir = os.path.join(tmp_out, "tables")
        assert os.path.isdir(tables_dir)
        files = os.listdir(tables_dir)
        assert any(f.endswith(".csv") for f in files)
        assert any(f.endswith(".html") for f in files)


class TestV16HierarchicalDocumentModel:
    def _assert_valid_tree(self, tmp_out):
        doc = json.load(open(os.path.join(tmp_out, "document.json"), encoding="utf-8"))
        assert doc["schema_version"] == "1.0"
        assert doc["root_element_id"] == "document-root"
        by_id = {el["id"]: el for el in doc["elements"]}
        assert len(by_id) == doc["element_count"]
        for el in doc["elements"]:
            assert "parent_id" in el and "child_ids" in el
            for child_id in el["child_ids"]:
                assert by_id[child_id]["parent_id"] == el["id"]
        return doc, by_id

    def test_xlsx_sheet_owns_table_element(self, xlsx_merged, tmp_out):
        _run(xlsx_merged, tmp_out)
        doc, by_id = self._assert_valid_tree(tmp_out)
        sheet = next(el for el in doc["elements"] if el["type"] == "sheet")
        assert any(by_id[c]["type"] == "table" for c in sheet["child_ids"])

    def test_pptx_slide_has_shape_level_children(self, pptx_merge_table, tmp_out):
        _run(pptx_merge_table, tmp_out)
        doc, by_id = self._assert_valid_tree(tmp_out)
        slide = next(el for el in doc["elements"] if el["type"] == "slide")
        child_types = {by_id[c]["type"] for c in slide["child_ids"]}
        assert "table" in child_types

    def test_docx_heading_parents_following_paragraph(self, tmp_path, tmp_out):
        from docx import Document
        docx = Document()
        docx.add_heading("Section", level=1)
        docx.add_paragraph("Paragraph under the section heading.")
        path = str(tmp_path / "hierarchy.docx")
        docx.save(path)
        _run(path, tmp_out)
        doc, by_id = self._assert_valid_tree(tmp_out)
        heading = next(el for el in doc["elements"] if el["type"] == "heading")
        assert any(by_id[c]["type"] == "paragraph" for c in heading["child_ids"])


class TestV16BoundedChunking:
    def test_oversized_single_element_is_split_below_limit(self, tmp_path, tmp_out):
        from docx import Document
        doc = Document()
        payload = "A" * 5100
        doc.add_paragraph(payload)
        path = str(tmp_path / "long.docx")
        doc.save(path)
        _run(path, tmp_out)
        chunks = [json.loads(line) for line in open(
            os.path.join(tmp_out, "chunks.jsonl"), encoding="utf-8")]
        assert len(chunks) == 3
        assert all(chunk["char_count"] <= 2000 for chunk in chunks)
        assert "".join(chunk["text"] for chunk in chunks) == payload
        assert [chunk["part_index"] for chunk in chunks] == [1, 2, 3]
        assert all(chunk["part_count"] == 3 for chunk in chunks)


class TestV16FailureIsolation:
    def test_failed_rerun_removes_stale_success_artifacts(self, xlsx_merged, tmp_path, tmp_out):
        first = _run(xlsx_merged, tmp_out)
        assert first["status"] == "passed"
        assert os.path.exists(os.path.join(tmp_out, "document.json"))
        assert os.path.isdir(os.path.join(tmp_out, "tables"))

        broken = tmp_path / "broken.xlsx"
        broken.write_bytes(b"PK-not-a-real-workbook")
        second = _run(str(broken), tmp_out)
        assert second["status"] == "failed"
        assert second["reason"] == "conversion_error"
        assert not os.path.exists(os.path.join(tmp_out, "document.json"))
        assert not os.path.exists(os.path.join(tmp_out, "chunks.jsonl"))
        assert not os.path.exists(os.path.join(tmp_out, "tables"))

    def test_missing_input_returns_machine_readable_failure(self, tmp_path, tmp_out):
        report = _run(str(tmp_path / "missing.pdf"), tmp_out)
        assert report["status"] == "failed"
        assert report["reason"] == "conversion_error"
        assert report["error_type"] == "FileNotFoundError"


class TestV16CanonicalContracts:
    def test_every_element_has_frozen_schema_fields(self, pptx_merge_table, tmp_out):
        _run(pptx_merge_table, tmp_out)
        document = json.load(open(os.path.join(tmp_out, "document.json"), encoding="utf-8"))
        required = {"id", "type", "parent_id", "children", "content", "content_format",
                    "heading_path", "engine", "confidence", "source_locator", "properties",
                    "warnings", "ordinal"}
        locator_keys = {"page", "slide", "sheet", "cell_range", "shape_id", "shape_name",
                        "bbox", "table_index", "relationship_id", "part", "source_file"}
        for element in document["elements"]:
            assert required <= element.keys()
            assert locator_keys <= element["source_locator"].keys()

    def test_table_assets_use_common_schema(self, xlsx_merged, tmp_out):
        _run(xlsx_merged, tmp_out)
        index = json.load(open(os.path.join(tmp_out, "tables", "index.json"), encoding="utf-8"))
        table = json.load(open(os.path.join(tmp_out, "tables", index[0]["assets"]["json"]),
                               encoding="utf-8"))
        assert table["schema_version"] == "1.0"
        assert table["dimensions"] == {"rows": 2, "columns": 2}
        assert table["source_format"] == "xlsx"
        assert table["source_locator"]["cell_range"] == "A1:B2"
        assert all("column" in cell for cell in table["cells"])

    def test_chunk_element_ids_and_locators_are_self_contained(self, pdf_short_digital, tmp_out):
        _run(pdf_short_digital, tmp_out)
        document = json.load(open(os.path.join(tmp_out, "document.json"), encoding="utf-8"))
        element_ids = {element["id"] for element in document["elements"]}
        chunks = [json.loads(line) for line in open(
            os.path.join(tmp_out, "chunks.jsonl"), encoding="utf-8")]
        assert chunks
        assert all(set(chunk["element_ids"]) <= element_ids for chunk in chunks)
        assert all(chunk["source_file"] == os.path.basename(pdf_short_digital) for chunk in chunks)
        assert all(chunk["page_start"] == 1 and chunk["page_end"] == 1 for chunk in chunks)


class TestV16FormatGranularity:
    def test_xlsx_splits_blocks_at_blank_rows(self, xlsx_two_blocks, tmp_out):
        _run(xlsx_two_blocks, tmp_out)
        document = json.load(open(os.path.join(tmp_out, "document.json"), encoding="utf-8"))
        tables = [element for element in document["elements"] if element["type"] == "table"]
        assert [table["source_locator"]["cell_range"] for table in tables] == ["A1:B2", "A5:B6"]

    def test_pdf_has_page_and_located_text_block(self, pdf_short_digital, tmp_out):
        _run(pdf_short_digital, tmp_out)
        document = json.load(open(os.path.join(tmp_out, "document.json"), encoding="utf-8"))
        page = next(element for element in document["elements"] if element["type"] == "page")
        child = next(element for element in document["elements"] if element["parent_id"] == page["id"])
        assert child["type"] in ("heading", "paragraph")
        assert child["source_locator"]["bbox"] is not None

    def test_pptx_picture_has_deterministic_asset_and_locator(self, pptx_picture, tmp_out):
        _run(pptx_picture, tmp_out)
        document = json.load(open(os.path.join(tmp_out, "document.json"), encoding="utf-8"))
        image = next(element for element in document["elements"] if element["type"] == "image")
        assert image["asset"].startswith("slide-0001-shape-")
        assert image["source_locator"]["relationship_id"].startswith("rId")
        assert image["source_locator"]["shape_id"] is not None
        assert os.path.isfile(os.path.join(tmp_out, "assets", image["asset"]))

    def test_group_shape_is_a_real_parent(self, pptx_group_with_table, tmp_out):
        _run(pptx_group_with_table, tmp_out)
        document = json.load(open(os.path.join(tmp_out, "document.json"), encoding="utf-8"))
        by_id = {element["id"]: element for element in document["elements"]}
        group = next(element for element in document["elements"] if element["type"] == "group")
        assert group["children"]
        assert any(by_id[child]["type"] == "table" for child in group["children"])

    def test_standalone_image_ocr_is_really_routed(self, image_with_mock_ocr, tmp_out):
        report = _run(image_with_mock_ocr, tmp_out)
        assert report["status"] == "passed"
        assert report["file_type"] == "image"
        md = open(os.path.join(tmp_out, "document.md"), encoding="utf-8").read()
        assert "Hello world" in md
        document = json.load(open(os.path.join(tmp_out, "document.json"), encoding="utf-8"))
        assert any(element["type"] == "ocr_region" for element in document["elements"])


class TestV16BundleValidator:
    def test_generated_bundle_self_validates(self, xlsx_merged, tmp_out):
        report = _run(xlsx_merged, tmp_out)
        assert report["bundle_validation"]["status"] == "passed"
        from validate_bundle import validate_bundle
        assert validate_bundle(tmp_out)["status"] == "passed"

    def test_validator_detects_broken_chunk_reference(self, xlsx_merged, tmp_out):
        _run(xlsx_merged, tmp_out)
        chunks_path = os.path.join(tmp_out, "chunks.jsonl")
        chunks = [json.loads(line) for line in open(chunks_path, encoding="utf-8")]
        chunks[0]["element_ids"] = ["missing-element"]
        with open(chunks_path, "w", encoding="utf-8") as f:
            for chunk in chunks:
                f.write(json.dumps(chunk) + "\n")
        from validate_bundle import validate_bundle
        result = validate_bundle(tmp_out)
        assert result["status"] == "failed"
        assert any("missing element" in error for error in result["errors"])

    def test_semantic_bundle_is_deterministic_across_reruns(self, xlsx_merged, tmp_path):
        out_a, out_b = str(tmp_path / "a"), str(tmp_path / "b")
        _run(xlsx_merged, out_a)
        _run(xlsx_merged, out_b)
        relative_files = ["document.md", "document.json", "chunks.jsonl",
                          "tables/index.json"]
        for relative in relative_files:
            assert open(os.path.join(out_a, relative), "rb").read() == open(
                os.path.join(out_b, relative), "rb").read()


class TestV16TableChunking:
    def test_large_pipe_table_repeats_header(self):
        from chunker import _split_markdown
        header = "| Name | Value |\n| --- | --- |"
        table = header + "\n" + "\n".join(
            f"| row-{index:04d} | {'x' * 30} |" for index in range(200))
        parts = _split_markdown(table, 500)
        assert len(parts) > 1
        assert all(len(part) <= 500 for part in parts)
        assert all(part.startswith(header) for part in parts)


class TestSecurityFixes:
    def test_eml_attachment_path_traversal_sanitized(self, tmp_path, tmp_out):
        import email.message
        msg = email.message.EmailMessage()
        msg["Subject"] = "test"
        msg["From"] = "a@example.com"
        msg["To"] = "b@example.com"
        msg.set_content("body text")
        msg.add_attachment(b"data", maintype="application", subtype="octet-stream",
                            filename="../../evil.bin")
        path = str(tmp_path / "test.eml")
        with open(path, "wb") as f:
            f.write(bytes(msg))
        report = _run(path, tmp_out)
        assert report["status"] == "passed"
        assets_dir = os.path.join(tmp_out, "assets")
        for fname in os.listdir(assets_dir):
            assert ".." not in fname
            assert not fname.startswith("/")
